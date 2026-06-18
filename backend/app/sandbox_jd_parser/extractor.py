# pyrefly: ignore [missing-import]
import docx

import re
from typing import Dict, Any, List

class JDExtractor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = docx.Document(file_path)

    def parse(self) -> Dict[str, Any]:
        """
        Parses the docx file to isolate target requirements using table parsing
        and token/regex paragraph matching.
        """
        # Initialize defaults
        extracted = {
            "job_title": "Unknown",
            "department": "Unknown",
            "experience_required": "Unknown",
            "core_technical_skills": [],
            "behavioral_competencies": [],
            "detailed_responsibilities": "Unknown"
        }

        # 1. Parse tables for key-value structural data
        for table in self.doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    key = row.cells[0].text.strip().lower()
                    val = row.cells[1].text.strip()
                    if not val:
                        continue
                    if "job title" in key or "title" == key:
                        extracted["job_title"] = val
                    elif "department" in key:
                        extracted["department"] = val
                    elif "experience" in key:
                        extracted["experience_required"] = val
                    elif "technical skills" in key or "core skills" in key or "technical requirements" in key:
                        # Split skills by standard dividers
                        items = [s.strip() for s in re.split(r'[,;\n•*-]', val) if s.strip()]
                        extracted["core_technical_skills"] = items
                    elif "behavioral" in key or "competencies" in key or "soft skills" in key:
                        items = [s.strip() for s in re.split(r'[,;\n•*-]', val) if s.strip()]
                        extracted["behavioral_competencies"] = items
                    elif "responsibilities" in key or "detailed responsibilities" in key:
                        extracted["detailed_responsibilities"] = val

        # 2. Parse paragraphs for keyword/regex patterns
        full_text = "\n".join([p.text.strip() for p in self.doc.paragraphs if p.text.strip()])
        lines = full_text.split('\n')
        
        current_section = None
        section_text = {
            "job_title": "",
            "department": "",
            "experience_required": "",
            "core_technical_skills": [],
            "behavioral_competencies": [],
            "detailed_responsibilities": []
        }

        for line in lines:
            line_lower = line.lower().strip()
            normalized_line = line_lower.rstrip(':').strip()
            
            # Check for header transitions
            
            # Job Title
            m_job_title = re.match(r'(?i)^(job\s+title|title)\s*:\s*(.*)', line)
            if not m_job_title and normalized_line in ["job title", "title"]:
                current_section = "job_title"
                continue
            elif m_job_title:
                current_section = "job_title"
                section_text["job_title"] = m_job_title.group(2).strip()
                continue
                
            # Department
            m_dept = re.match(r'(?i)^(department)\s*:\s*(.*)', line)
            if not m_dept and normalized_line in ["department"]:
                current_section = "department"
                continue
            elif m_dept:
                current_section = "department"
                section_text["department"] = m_dept.group(2).strip()
                continue

            # Experience
            m_exp = re.match(r'(?i)^(experience\s+required|experience(\s+level)?)\s*:\s*(.*)', line)
            if not m_exp and normalized_line in ["experience required", "experience", "experience level"]:
                current_section = "experience_required"
                continue
            elif m_exp:
                current_section = "experience_required"
                section_text["experience_required"] = m_exp.group(3).strip()
                continue

            # Core Technical Skills
            m_skills = re.match(r'(?i)^((core\s+)?technical\s+(skills?|requirements?)|skills?)\s*:\s*(.*)', line)
            if not m_skills and normalized_line in ["core technical skills", "technical skills", "skills", "technical requirements"]:
                current_section = "core_technical_skills"
                continue
            elif m_skills:
                current_section = "core_technical_skills"
                if m_skills.group(4).strip():
                    section_text["core_technical_skills"].append(m_skills.group(4).strip())
                continue

            # Behavioral Competencies
            m_beh = re.match(r'(?i)^(behavioral\s+(competencies|skills?)|soft\s+skills?)\s*:\s*(.*)', line)
            if not m_beh and normalized_line in ["behavioral competencies", "behavioral skills", "soft skills"]:
                current_section = "behavioral_competencies"
                continue
            elif m_beh:
                current_section = "behavioral_competencies"
                if m_beh.group(3).strip():
                    section_text["behavioral_competencies"].append(m_beh.group(3).strip())
                continue

            # Detailed Responsibilities
            m_resp = re.match(r'(?i)^((detailed\s+|key\s+)?responsibilities)\s*:\s*(.*)', line)
            if not m_resp and normalized_line in ["detailed responsibilities", "responsibilities", "key responsibilities"]:
                current_section = "detailed_responsibilities"
                continue
            elif m_resp:
                current_section = "detailed_responsibilities"
                if m_resp.group(3).strip():
                    section_text["detailed_responsibilities"].append(m_resp.group(3).strip())
                continue
                
            # Accumulate values under active sections
            if current_section:
                if current_section in ["job_title", "department", "experience_required"]:
                    if not section_text[current_section]:
                        section_text[current_section] = line.strip()
                    else:
                        section_text[current_section] += " " + line.strip()
                elif current_section in ["core_technical_skills", "behavioral_competencies"]:
                    # Clean out list styling symbols
                    cleaned = re.sub(r'^[-*•\d\.\s]+', '', line).strip()
                    if cleaned:
                        items = [s.strip() for s in re.split(r'[,;]', cleaned) if s.strip()]
                        section_text[current_section].extend(items)
                elif current_section == "detailed_responsibilities":
                    section_text["detailed_responsibilities"].append(line.strip())

        # Merge paragraph-extracted values if they were missing or not set by tables
        if extracted["job_title"] == "Unknown" and section_text["job_title"]:
            extracted["job_title"] = section_text["job_title"]
        if extracted["department"] == "Unknown" and section_text["department"]:
            extracted["department"] = section_text["department"]
        if extracted["experience_required"] == "Unknown" and section_text["experience_required"]:
            extracted["experience_required"] = section_text["experience_required"]
        if not extracted["core_technical_skills"] and section_text["core_technical_skills"]:
            extracted["core_technical_skills"] = [s for s in section_text["core_technical_skills"] if s]
        if not extracted["behavioral_competencies"] and section_text["behavioral_competencies"]:
            extracted["behavioral_competencies"] = [s for s in section_text["behavioral_competencies"] if s]
        if extracted["detailed_responsibilities"] == "Unknown" and section_text["detailed_responsibilities"]:
            extracted["detailed_responsibilities"] = "\n".join(section_text["detailed_responsibilities"]).strip()

        return extracted
